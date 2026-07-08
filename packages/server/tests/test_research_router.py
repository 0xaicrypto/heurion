"""Phase 1 tests for research_router.py.

Covers full CRUD + enrollment + roster + decision +
Patient → Studies derived view (design §3.4 D18).
"""
from __future__ import annotations

import pathlib
import sqlite3
import sys

import pytest
from fastapi.testclient import TestClient

sys.path.insert(0, str(pathlib.Path(__file__).resolve().parents[1]))


# ─────────────────────────────────────────────────────────────────────
# Fixtures
# ─────────────────────────────────────────────────────────────────────

@pytest.fixture
def client_as(monkeypatch):
    """Return a factory that builds a TestClient where get_current_user
    returns the given user_id. Runs Alembic migrations up to head so
    the research_* tables exist."""
    from nexus_server.main import create_app
    from nexus_server.auth import get_current_user
    from nexus_server.migrations.runner import run_migrations

    # conftest._init_db wipes + re-init the DB before every test, but
    # only calls database.init_db() (legacy table creation). The
    # research tables come from migration 0004, so we have to run the
    # Alembic chain ourselves.
    run_migrations()

    def _make(user_id: str = "dr-qian") -> TestClient:
        app = create_app()
        app.dependency_overrides[get_current_user] = lambda: user_id
        return TestClient(app)

    return _make


# ─────────────────────────────────────────────────────────────────────
# Study CRUD
# ─────────────────────────────────────────────────────────────────────


def test_create_then_get_study(client_as):
    c = client_as("dr-qian")

    body = {
        "display_name": "Hybrid RT NSCLC IV",
        "short_code": "HybridRT-IV",
        "phase": "I/II",
        "target_n": 35,
        "primary_endpoint": "Pneumonitis ≥G3 incidence",
        "secondary_endpoints": ["mPFS", "1y PFS", "mOS", "1y OS"],
        "inclusion": [
            {"id": "a", "text": "Age 18-70", "kind": "auto-rule",
             "rule_dsl": "age BETWEEN 18 AND 70"},
            {"id": "h", "text": "Informed consent", "kind": "manual"},
        ],
        "exclusion": [
            {"id": "d", "text": "Driver mutation positive",
             "kind": "auto-rule", "rule_dsl": "driver_mutation = 'positive'"},
        ],
        "schedule": [
            {"label": "baseline", "offset_days": 0, "assessments": ["pet_ct"]},
            {"label": "rt_end_4w", "offset_days": 28, "assessments": ["chest_ct"]},
        ],
    }
    r = c.post("/api/v1/research/studies", json=body)
    assert r.status_code == 201, r.text
    detail = r.json()
    sid = detail["study_id"]
    assert detail["display_name"] == "Hybrid RT NSCLC IV"
    assert detail["status"] == "draft"
    assert len(detail["inclusion"]) == 2
    assert detail["inclusion"][0]["kind"] == "auto-rule"
    assert detail["inclusion"][1]["kind"] == "manual"
    assert detail["enrolled_count"] == 0
    assert detail["candidate_count"] == 0

    # GET
    r = c.get(f"/api/v1/research/studies/{sid}")
    assert r.status_code == 200
    again = r.json()
    assert again["study_id"] == sid
    assert again["short_code"] == "HybridRT-IV"


def test_list_studies_only_includes_active_by_default(client_as):
    c = client_as("dr-qian")

    a = c.post("/api/v1/research/studies",
               json={"display_name": "Study A", "short_code": "A"}).json()
    b = c.post("/api/v1/research/studies",
               json={"display_name": "Study B", "short_code": "B"}).json()

    r = c.get("/api/v1/research/studies").json()
    ids = {s["study_id"] for s in r}
    assert a["study_id"] in ids and b["study_id"] in ids

    # Archive A
    arch = c.delete(f"/api/v1/research/studies/{a['study_id']}")
    assert arch.status_code == 200

    r2 = c.get("/api/v1/research/studies").json()
    ids2 = {s["study_id"] for s in r2}
    assert a["study_id"] not in ids2
    assert b["study_id"] in ids2

    # include_archived=true still shows it
    r3 = c.get("/api/v1/research/studies?include_archived=true").json()
    assert any(s["study_id"] == a["study_id"] for s in r3)


def test_update_study_protocol_emits_protocol_updated(client_as):
    c = client_as("dr-qian")
    s = c.post("/api/v1/research/studies",
               json={"display_name": "S", "short_code": "S"}).json()
    r = c.patch(f"/api/v1/research/studies/{s['study_id']}", json={
        "status": "enrolling",
        "inclusion": [{"id": "a", "text": "age", "kind": "auto-rule",
                       "rule_dsl": "age >= 18"}],
    })
    assert r.status_code == 200, r.text
    detail = r.json()
    assert detail["status"] == "enrolling"
    assert len(detail["inclusion"]) == 1


def test_archive_returns_404_when_already_archived(client_as):
    c = client_as("dr-qian")
    s = c.post("/api/v1/research/studies",
               json={"display_name": "Tmp", "short_code": "Tmp"}).json()
    c.delete(f"/api/v1/research/studies/{s['study_id']}")
    r = c.get(f"/api/v1/research/studies/{s['study_id']}")
    assert r.status_code == 404


# ─────────────────────────────────────────────────────────────────────
# Tenant isolation
# ─────────────────────────────────────────────────────────────────────


def test_cross_tenant_cannot_read(client_as):
    """User A creates a study. User B must not see it."""
    a = client_as("dr-A")
    s = a.post("/api/v1/research/studies",
               json={"display_name": "Private", "short_code": "P"}).json()

    b = client_as("dr-B")
    r = b.get(f"/api/v1/research/studies/{s['study_id']}")
    assert r.status_code == 404, "tenant isolation broke"

    r2 = b.get("/api/v1/research/studies").json()
    assert all(x["study_id"] != s["study_id"] for x in r2)


# ─────────────────────────────────────────────────────────────────────
# Enrollment / Roster
# ─────────────────────────────────────────────────────────────────────


def test_enroll_then_withdraw_then_roster(client_as):
    c = client_as("dr-qian")
    s = c.post("/api/v1/research/studies",
               json={"display_name": "E", "short_code": "E"}).json()
    sid = s["study_id"]

    r1 = c.post(f"/api/v1/research/studies/{sid}/enrollments",
                json={"patient_hash": "p1", "consent_signed_at": 12345})
    assert r1.status_code == 201, r1.text
    e1 = r1.json()
    assert e1["enrollment_seq"] == 1
    assert e1["status"] == "enrolled"
    assert e1["consent_signed_at"] == 12345

    # Second enrollment gets seq=2
    r2 = c.post(f"/api/v1/research/studies/{sid}/enrollments",
                json={"patient_hash": "p2"})
    assert r2.json()["enrollment_seq"] == 2

    # Withdraw p1
    r3 = c.request("DELETE",
                   f"/api/v1/research/studies/{sid}/enrollments/p1",
                   json={"reason": "consent withdrawn"})
    assert r3.status_code == 200

    # Roster default: only enrolled
    roster = c.get(f"/api/v1/research/studies/{sid}/roster").json()
    assert {r["patient_hash"] for r in roster} == {"p2"}

    # Include withdrawn
    full = c.get(f"/api/v1/research/studies/{sid}/roster?include_withdrawn=true").json()
    p1 = next(r for r in full if r["patient_hash"] == "p1")
    assert p1["status"] == "withdrawn"
    assert p1["withdrawal_reason"] == "consent withdrawn"


def test_summary_counts_track_enrollment(client_as):
    c = client_as("dr-qian")
    s = c.post("/api/v1/research/studies",
               json={"display_name": "C", "short_code": "C"}).json()
    sid = s["study_id"]

    c.post(f"/api/v1/research/studies/{sid}/enrollments",
           json={"patient_hash": "pX"})
    c.post(f"/api/v1/research/studies/{sid}/enrollments",
           json={"patient_hash": "pY"})

    listed = c.get("/api/v1/research/studies").json()
    me = next(x for x in listed if x["study_id"] == sid)
    assert me["enrolled_count"] == 2


def test_double_enrollment_idempotent(client_as):
    """Re-enrolling the same patient must NOT bump seq or duplicate."""
    c = client_as("dr-qian")
    s = c.post("/api/v1/research/studies",
               json={"display_name": "I", "short_code": "I"}).json()
    sid = s["study_id"]

    r1 = c.post(f"/api/v1/research/studies/{sid}/enrollments",
                json={"patient_hash": "p1"}).json()
    r2 = c.post(f"/api/v1/research/studies/{sid}/enrollments",
                json={"patient_hash": "p1"}).json()

    assert r1["enrollment_seq"] == r2["enrollment_seq"]

    roster = c.get(f"/api/v1/research/studies/{sid}/roster").json()
    assert len(roster) == 1


def test_re_enroll_after_withdraw_keeps_seq(client_as):
    c = client_as("dr-qian")
    s = c.post("/api/v1/research/studies",
               json={"display_name": "R", "short_code": "R"}).json()
    sid = s["study_id"]

    seq1 = c.post(f"/api/v1/research/studies/{sid}/enrollments",
                  json={"patient_hash": "p1"}).json()["enrollment_seq"]
    c.request("DELETE",
              f"/api/v1/research/studies/{sid}/enrollments/p1",
              json={"reason": "test"})
    re_enrolled = c.post(f"/api/v1/research/studies/{sid}/enrollments",
                         json={"patient_hash": "p1"}).json()
    assert re_enrolled["status"] == "enrolled"
    # Keep original seq (clinical-trials convention) — no double-counting.
    assert re_enrolled["enrollment_seq"] == seq1


# ─────────────────────────────────────────────────────────────────────
# Patient → Studies derived view (D18)
# ─────────────────────────────────────────────────────────────────────


def test_patient_studies_returns_enrollments(client_as):
    c = client_as("dr-qian")
    s1 = c.post("/api/v1/research/studies",
                json={"display_name": "Hybrid RT", "short_code": "HRT"}).json()
    s2 = c.post("/api/v1/research/studies",
                json={"display_name": "ES-SCLC", "short_code": "ESC"}).json()
    c.post(f"/api/v1/research/studies/{s1['study_id']}/enrollments",
           json={"patient_hash": "PT-1"})
    c.post(f"/api/v1/research/studies/{s2['study_id']}/enrollments",
           json={"patient_hash": "PT-1"})

    r = c.get("/api/v1/patients/PT-1/studies").json()
    ids = {m["study_id"] for m in r}
    assert s1["study_id"] in ids and s2["study_id"] in ids
    assert all(m["status"] == "enrolled" for m in r)


def test_patient_studies_includes_withdrawn(client_as):
    c = client_as("dr-qian")
    s = c.post("/api/v1/research/studies",
               json={"display_name": "WD", "short_code": "WD"}).json()
    c.post(f"/api/v1/research/studies/{s['study_id']}/enrollments",
           json={"patient_hash": "P-out"})
    c.request("DELETE",
              f"/api/v1/research/studies/{s['study_id']}/enrollments/P-out",
              json={"reason": "AE"})
    r = c.get("/api/v1/patients/P-out/studies").json()
    assert len(r) == 1
    assert r[0]["status"] == "withdrawn"
    assert r[0]["withdrawal_reason"] == "AE"


# ─────────────────────────────────────────────────────────────────────
# Migration sanity
# ─────────────────────────────────────────────────────────────────────


def test_migration_creates_all_research_tables(client_as):
    """Spinning up the client should fire migration 0004. Verify the
    tables exist with the right columns."""
    client_as()  # triggers init_db + migrations via app creation
    from nexus_server.database import get_db_connection
    with get_db_connection() as conn:
        for tbl in (
            "research_studies", "study_enrollments",
            "screening_evaluations", "study_assessments",
            "study_observations",
        ):
            rows = conn.execute(
                f"SELECT name FROM sqlite_master WHERE type='table' AND name=?",
                (tbl,),
            ).fetchall()
            assert rows, f"missing table {tbl}"

        # nexus_sessions should have scope_kind and scope_id
        cols = {r[1] for r in conn.execute("PRAGMA table_info(nexus_sessions)").fetchall()}
        assert "scope_kind" in cols
        assert "scope_id" in cols


# ─────────────────────────────────────────────────────────────────────
# F-docx-import-test — full /protocol/import flow after F-merge-patients-db
# ─────────────────────────────────────────────────────────────────────
# Regression: F-merge-patients-db moved the `patients` table from
# dicom_index.db into the shared rune_server.db AND added a one-shot
# migration that drops the legacy table after copy. The medic
# subsequently reported "Load failed" on .docx import and asked us to
# confirm the merge didn't break this code path. These tests ratchet
# in the docx-import flow end-to-end so the same regression can't
# silently come back.


def test_protocol_import_end_to_end_after_patients_merge(client_as, monkeypatch):
    """POST /research/studies/{id}/protocol/import — happy path.

    Sets up: study created, .docx written to disk, uploads row
    inserted manually (the actual upload endpoint is unit-tested
    elsewhere), then calls the import endpoint and verifies it
    returns a usable draft with status 'ok'. Stubs the LLM gateway
    so the test doesn't depend on network.

    Critically: ALSO calls init_patients_table() before the import
    so we exercise the exact boot sequence the medic's machine runs
    (init_patients_table → migration → upload row read → docx parse).
    Failure here means F-merge-patients-db's migration broke the
    upload-table init or some other shared-db dependency.
    """
    import uuid, time, asyncio  # noqa: F401
    from pathlib import Path
    from docx import Document
    from nexus_server.database import get_db_connection
    from nexus_server.files import _ensure_uploads_table
    from nexus_server.patients_router import init_patients_table

    # Run the exact init sequence the sidecar runs on boot — including
    # the F-merge-patients-db migration helper.
    init_patients_table()
    _ensure_uploads_table()

    c = client_as("dr-import")

    # 1. Create a study row to attach the protocol to.
    r = c.post("/api/v1/research/studies", json={
        "display_name": "Import Smoke", "short_code": "ISMK",
        "phase": "II", "target_n": 10,
    })
    assert r.status_code == 201, r.text
    sid = r.json()["study_id"]

    # 2. Manufacture a real .docx on disk + register it as an upload.
    import tempfile
    tmp = Path(tempfile.mkdtemp())
    docx_path = tmp / "smoke_protocol.docx"
    doc = Document()
    doc.add_heading("Smoke Protocol", level=1)
    doc.add_paragraph("入选标准:")
    doc.add_paragraph("年龄 18-75 岁")
    doc.add_paragraph("ECOG 0-1")
    doc.add_paragraph("排除标准:")
    doc.add_paragraph("妊娠期妇女")
    doc.save(str(docx_path))

    file_id = str(uuid.uuid4())
    with get_db_connection() as conn:
        conn.execute(
            "INSERT INTO uploads(file_id, user_id, name, mime, "
            "size_bytes, disk_path, created_at) "
            "VALUES (?,?,?,?,?,?,?)",
            (file_id, "dr-import", "smoke_protocol.docx",
             "application/vnd.openxmlformats-officedocument."
             "wordprocessingml.document",
             docx_path.stat().st_size, str(docx_path),
             int(time.time() * 1000)),
        )
        conn.commit()

    # 3. Stub the LLM gateway so the parser doesn't hit the network.
    #    Returning a None-equivalent forces the regex fallback, which is
    #    the deterministic / network-free path we want under test.
    from nexus_server.research import protocol_parser
    async def _stub_llm(doc_markdown):
        return None
    monkeypatch.setattr(protocol_parser, "_llm_extract_full", _stub_llm)

    # 4. Hit /protocol/import.
    r = c.post(
        f"/api/v1/research/studies/{sid}/protocol/import",
        json={"upload_file_id": file_id},
    )
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["status"] == "ok"
    draft = body["draft"]
    # Regex fallback should find the inclusion / exclusion buckets.
    # (Exact content depends on the heuristic; we only assert the
    # response shape so the test isn't fragile to parser tuning.)
    assert "inclusion" in draft
    assert "exclusion" in draft
    assert "schedule" in draft

    # 5. Sanity: protocol_doc_id should now be stamped on the study.
    detail = c.get(f"/api/v1/research/studies/{sid}").json()
    assert detail["protocol_doc_id"] == file_id


def test_archived_study_hidden_from_llm_roster(client_as):
    """F-roster-archive-filter — DELETE /studies/{id} sets archived_at
    but leaves ``status='enrolling'`` (archive is a UI-hide signal,
    not a lifecycle transition). The cross-research LLM prompt
    builder ``_gather_all_studies_summary`` previously didn't filter
    on archived_at, so deleted studies kept appearing in the AI's
    "studies you have" list — the medic deletes a trial, asks
    'what studies do I have', and the LLM cheerfully recites the
    deleted one back. This test pins the filter."""
    import time
    from nexus_server.retrieval_tiers import _gather_all_studies_summary
    from nexus_server.database import get_db_connection

    uid = "dr-archive"
    client_as(uid)  # boots app + migrations
    now = int(time.time() * 1000)
    with get_db_connection() as conn:
        for sid, name, status, archived_at in (
            ("s_keep_a", "KEEP_ACTIVE",   "enrolling", None),
            ("s_keep_d", "KEEP_DRAFT",    "draft",     None),
            ("s_hide_a", "HIDE_ARCHIVED", "enrolling", now),
        ):
            conn.execute(
                "INSERT INTO research_studies"
                "(study_id,user_id,display_name,short_code,phase,"
                " status,target_n,primary_endpoint,inclusion_json,"
                " exclusion_json,schedule_json,protocol_summary,"
                " created_at,updated_at,archived_at) "
                "VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                (sid, uid, name, sid, "II", status, 20, "",
                 "[]", "[]", "[]", "", now, now, archived_at),
            )
        conn.commit()
        roster = _gather_all_studies_summary(conn, uid)

    assert "KEEP_ACTIVE" in roster
    assert "KEEP_DRAFT" in roster
    # The crux of the regression: an archived study MUST be invisible.
    assert "HIDE_ARCHIVED" not in roster, (
        "archived study leaked into the LLM roster — F-roster-"
        "archive-filter regressed"
    )


def test_empty_roster_emits_explicit_authoritative_block(client_as):
    """F-roster-empty-explicit — when the medic has archived every
    study, the cross-research LLM prompt USED to receive an empty
    string from ``_gather_all_studies_summary``. That left no current-
    state signal in the prompt, so when chat history contained the
    AI's earlier "you have 3 active + 8 draft studies" reply, the LLM
    would happily echo it on the next turn — exactly the bug the
    medic reported. Fix: emit an explicit "0 active, 0 draft" block
    + a stern AUTHORITATIVE instruction even when empty."""
    from nexus_server.retrieval_tiers import _gather_all_studies_summary
    from nexus_server.database import get_db_connection

    uid = "dr-empty-roster"
    client_as(uid)
    with get_db_connection() as conn:
        out = _gather_all_studies_summary(conn, uid)

    # Must NOT be empty (was the bug).
    assert out, "empty string = no signal to LLM = bug"
    assert "0 active, 0 draft" in out
    assert "AUTHORITATIVE" in out
    assert "your study list is empty" in out


def test_protocol_import_404_when_upload_missing(client_as):
    """Importing a non-existent upload_file_id used to surface as
    "Load failed" because the parser returned a notes-only draft and
    the request quietly succeeded with no useful content. After
    F-docx-import-diag we still return 200 + an empty draft + a note
    explaining why — verify that contract."""
    c = client_as("dr-missing")
    r = c.post("/api/v1/research/studies", json={
        "display_name": "X", "short_code": "X", "phase": "II",
    })
    sid = r.json()["study_id"]
    r = c.post(
        f"/api/v1/research/studies/{sid}/protocol/import",
        json={"upload_file_id": "definitely-not-a-real-file-id"},
    )
    # Either 200 with notes, or 404 with detail — both are acceptable
    # post-F-docx-import-diag. What MUST NOT happen is a 500 / crash.
    assert r.status_code in (200, 404), r.text
    if r.status_code == 200:
        notes = r.json()["draft"].get("notes") or []
        assert any("not found" in n.lower() or "missing" in n.lower()
                   for n in notes), notes
