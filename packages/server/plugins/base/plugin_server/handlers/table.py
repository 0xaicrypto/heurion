import io
import os
import uuid
from typing import Any

from docx import Document

from ..storage import upload_output


def _make_file_id() -> str:
    return f"table_{uuid.uuid4().hex[:16]}"


def render_table(payload: dict[str, Any]) -> dict:
    """Render a medical/Table 1 baseline characteristics table as DOCX."""
    title = payload.get("title", "Table 1")
    headers = payload.get("headers", ["Variable", "Value"])
    rows = payload.get("rows", [])
    output_name = payload.get("output_name") or "table_1"

    doc = Document()
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)

    for row in rows:
        cells = table.add_row().cells
        for i, cell in enumerate(row):
            if i < len(cells):
                cells[i].text = str(cell)

    file_id = _make_file_id()
    file_name = f"{output_name}.docx"
    buf = io.BytesIO()
    doc.save(buf)

    tenant_prefix = os.environ.get("TENANT_PREFIX", "default/anonymous")
    return upload_output(
        file_obj=buf,
        file_name=file_name,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        tenant_prefix=tenant_prefix,
        file_id=file_id,
    )
